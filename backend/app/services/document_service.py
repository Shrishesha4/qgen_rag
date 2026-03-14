"""
Document service for file upload and processing.
"""

import os
import hashlib
import asyncio
from datetime import datetime, timezone
from typing import Optional, List, BinaryIO, Dict, Any
from pathlib import Path
import uuid
import aiofiles

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.document import Document, DocumentChunk
from app.models.question import Question
from app.schemas.document import DocumentResponse, DocumentListItem, PaginationInfo
from app.core.config import settings
from app.services.embedding_service import EmbeddingService


class DocumentService:
    """Service for document management and processing."""

    def __init__(self, db: AsyncSession, embedding_service: Optional[EmbeddingService] = None):
        self.db = db
        self.embedding_service = embedding_service or EmbeddingService()

    async def upload_document(
        self,
        user_id: uuid.UUID,
        filename: str,
        file_content: bytes,
        mime_type: str,
        index_type: str = "primary",
        subject_id: Optional[uuid.UUID] = None,
    ) -> Document:
        """Upload and process a document."""
        # Calculate file hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_size = len(file_content)

        # Check for duplicate (same user, same file)
        existing = await self.db.execute(
            select(Document).where(
                Document.user_id == user_id,
                Document.file_hash == file_hash,
            )
        )
        if existing.scalar_one_or_none():
            raise ValueError("Document already uploaded")

        # Create storage path
        storage_dir = Path(settings.UPLOAD_DIR) / str(user_id)
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        unique_filename = f"{uuid.uuid4()}_{filename}"
        storage_path = storage_dir / unique_filename

        # Save file
        async with aiofiles.open(storage_path, "wb") as f:
            await f.write(file_content)

        # Create document record
        document = Document(
            user_id=user_id,
            filename=filename,
            file_hash=file_hash,
            file_size_bytes=file_size,
            mime_type=mime_type,
            storage_path=str(storage_path),
            processing_status="pending",
            index_type=index_type,
            subject_id=subject_id,
        )
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        # Start async processing
        asyncio.create_task(self._process_document(document.id))

        return document

    async def upload_reference_document(
        self,
        user_id: uuid.UUID,
        filename: str,
        file_content: bytes,
        mime_type: str,
        subject_id: uuid.UUID,
        index_type: str,  # reference_book or template_paper
    ) -> Document:
        """
        Upload a reference document (book, template paper, or reference questions).
        These are stored in a separate index for novelty comparison.
        
        Unlike upload_document(), this does NOT auto-start _process_document.
        Processing is handled by the calling endpoint based on index_type.
        """
        if index_type not in ("primary", "reference_book", "template_paper", "reference_questions"):
            raise ValueError("index_type must be 'primary', 'reference_book', 'template_paper', or 'reference_questions'")
        
        # Calculate file hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_size = len(file_content)

        # If this user already linked the same file for this subject+type, return existing link.
        existing_link = await self.db.execute(
            select(Document).where(
                Document.user_id == user_id,
                Document.subject_id == subject_id,
                Document.index_type == index_type,
                Document.file_hash == file_hash,
            )
        )
        existing_doc = existing_link.scalar_one_or_none()
        if existing_doc:
            return existing_doc

        # Reuse any already-processed matching document from any user by creating a new
        # subject/user-scoped document record and cloning chunk rows (no re-embedding).
        reusable = await self.db.execute(
            select(Document).where(
                Document.file_hash == file_hash,
                Document.index_type == index_type,
                Document.processing_status == "completed",
            )
            .order_by(Document.processed_at.desc().nullslast())
        )
        source_doc = reusable.scalars().first()
        if source_doc:
            metadata = dict(source_doc.document_metadata or {})
            metadata["linked_from_document_id"] = str(source_doc.id)
            document = Document(
                user_id=user_id,
                filename=filename,
                file_hash=file_hash,
                file_size_bytes=file_size,
                mime_type=mime_type,
                storage_path=source_doc.storage_path,
                processing_status="completed",
                index_type=index_type,
                subject_id=subject_id,
                total_chunks=source_doc.total_chunks,
                total_tokens=source_doc.total_tokens,
                processed_at=source_doc.processed_at,
                document_metadata=metadata,
            )
            self.db.add(document)
            await self.db.flush()

            if index_type in ("reference_book", "template_paper"):
                source_chunks = await self.db.execute(
                    select(DocumentChunk)
                    .where(DocumentChunk.document_id == source_doc.id)
                    .order_by(DocumentChunk.chunk_index)
                )
                new_chunks = []
                for chunk in source_chunks.scalars().all():
                    new_chunks.append(
                        DocumentChunk(
                            document_id=document.id,
                            chunk_index=chunk.chunk_index,
                            chunk_text=chunk.chunk_text,
                            chunk_embedding=chunk.chunk_embedding,
                            token_count=chunk.token_count,
                            page_number=chunk.page_number,
                            section_heading=chunk.section_heading,
                            chunk_metadata=dict(chunk.chunk_metadata or {}),
                        )
                    )
                if new_chunks:
                    self.db.add_all(new_chunks)

            await self.db.commit()
            await self.db.refresh(document)
            return document

        # Create storage path
        storage_dir = Path(settings.UPLOAD_DIR) / str(user_id)
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        unique_filename = f"{uuid.uuid4()}_{filename}"
        storage_path = storage_dir / unique_filename

        # Save file
        async with aiofiles.open(storage_path, "wb") as f:
            await f.write(file_content)

        # Create document record
        document = Document(
            user_id=user_id,
            filename=filename,
            file_hash=file_hash,
            file_size_bytes=file_size,
            mime_type=mime_type,
            storage_path=str(storage_path),
            processing_status="pending",
            index_type=index_type,
            subject_id=subject_id,
        )
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        return document

    async def upload_and_process_document(
        self,
        user_id: uuid.UUID,
        filename: str,
        file_content: bytes,
        mime_type: str,
        context: Optional[str] = None,
        subject_id: Optional[uuid.UUID] = None,
    ) -> Document:
        """
        Upload and synchronously process a document.
        Used for quick generation where we need the document ready immediately.
        """
        # Calculate file hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_size = len(file_content)

        # Check for duplicate (same user, same file)
        existing = await self.db.execute(
            select(Document).where(
                Document.user_id == user_id,
                Document.file_hash == file_hash,
            )
        )
        existing_doc = existing.scalar_one_or_none()
        if existing_doc:
            # If already processed, return it
            if existing_doc.processing_status == "completed":
                return existing_doc
            raise ValueError("Document already uploaded and is processing")

        # Create storage path
        storage_dir = Path(settings.UPLOAD_DIR) / str(user_id)
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        unique_filename = f"{uuid.uuid4()}_{filename}"
        storage_path = storage_dir / unique_filename

        # Save file
        async with aiofiles.open(storage_path, "wb") as f:
            await f.write(file_content)

        # Create document record with context
        document = Document(
            user_id=user_id,
            filename=filename,
            file_hash=file_hash,
            file_size_bytes=file_size,
            mime_type=mime_type,
            storage_path=str(storage_path),
            processing_status="processing",
            subject_id=subject_id,
            document_metadata={"context": context} if context else {},
        )
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        try:
            # Extract text with page-level information (with OCR fallback)
            pages_data = await self._extract_text_with_pages(str(storage_path), mime_type)
            
            # Determine extraction method used
            extraction_methods = set(p.get("extraction_method", "native") for p in pages_data)
            used_ocr = "ocr" in extraction_methods
            
            # Update document metadata with page info
            document.document_metadata = {
                **(document.document_metadata or {}),
                "total_pages": len(pages_data),
                "document_name": pages_data[0].get("document_name") if pages_data else filename,
                "extraction_method": "ocr" if used_ocr else "native",
                "used_ocr": used_ocr,
            }

            # Chunk the document with page tracking
            chunks = self._chunk_text_with_pages(pages_data)

            # PERFORMANCE OPTIMIZATION: Batch generate all embeddings at once
            chunk_texts = [chunk_data["text"] for chunk_data in chunks]
            embeddings = await self.embedding_service.get_embeddings(chunk_texts, is_query=False)

            # Create chunk objects (batch insert optimization)
            total_tokens = 0
            chunk_objects = []
            for idx, (chunk_data, embedding) in enumerate(zip(chunks, embeddings)):
                # Build comprehensive chunk metadata
                chunk_metadata = chunk_data.get("metadata", {})
                chunk_metadata["source_info"] = {
                    "document_name": chunk_data.get("document_name", filename),
                    "page_number": chunk_data.get("page_number"),
                    "page_range": chunk_data.get("page_range"),
                    "position_in_page": chunk_data.get("position_in_page"),
                    "position_percentage": chunk_data.get("position_percentage"),
                }
                
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx,
                    chunk_text=chunk_data["text"],
                    chunk_embedding=embedding,
                    token_count=chunk_data["token_count"],
                    page_number=chunk_data.get("page_number"),
                    section_heading=chunk_data.get("section_heading"),
                    chunk_metadata=chunk_metadata,
                )
                chunk_objects.append(chunk)
                total_tokens += chunk_data["token_count"]
            
            # Bulk insert all chunks at once
            self.db.add_all(chunk_objects)

            # Update document
            document.total_chunks = len(chunks)
            document.total_tokens = total_tokens
            document.processing_status = "completed"
            document.processed_at = datetime.now(timezone.utc)
            await self.db.commit()
            await self.db.refresh(document)
            
            return document

        except Exception as e:
            document.processing_status = "failed"
            document.document_metadata = {"error": str(e), "context": context}
            await self.db.commit()
            raise ValueError(f"Document processing failed: {str(e)}")

    async def _update_progress(self, db, document, step: str, progress: int, detail: str = ""):
        """Update document processing progress metadata."""
        import logging
        logger = logging.getLogger(__name__)
        document.document_metadata = {
            **(document.document_metadata or {}),
            "processing_step": step,
            "processing_progress": progress,
            "processing_detail": detail,
        }
        await db.commit()
        logger.info(f"[{document.filename}] {step}: {progress}% - {detail}")

    async def _process_document(self, document_id: uuid.UUID):
        """Process document: extract text with page info, chunk, and create embeddings.
        
        Uses batched embedding generation and progress tracking for large files.
        Splits into phases to avoid holding DB connections during long extraction/OCR.
        """
        import logging
        logger = logging.getLogger(__name__)
        from app.core.database import AsyncSessionLocal
        from sqlalchemy import update as sa_update
        
        EMBEDDING_BATCH_SIZE = 32  # Process embeddings in batches to avoid OOM
        DB_INSERT_BATCH_SIZE = 100  # Insert chunks in batches to avoid huge transactions
        
        # Helper for short-lived progress updates (no long-held session)
        async def _quick_progress(step, progress, detail):
            async with AsyncSessionLocal() as s:
                await s.execute(
                    sa_update(Document)
                    .where(Document.id == document_id)
                    .values(document_metadata={
                        "processing_step": step,
                        "processing_progress": progress,
                        "processing_detail": detail,
                    })
                )
                await s.commit()
            logger.info(f"[doc:{document_id}] {step}: {progress}% - {detail}")
        
        # ── Phase 1: Fetch document info and mark as processing ──
        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            document = result.scalar_one_or_none()
            if not document:
                return
            
            storage_path = document.storage_path
            mime_type = document.mime_type
            doc_filename = document.filename
            
            document.processing_status = "processing"
            await self._update_progress(db, document, "extracting", 5, "Extracting text from document...")
        # Session released — no DB connection held during extraction
        
        try:
            # ── Phase 2: Extract text (long-running, especially with OCR) ──
            # Progress callback uses short-lived sessions
            pages_data = await self._extract_text_with_pages(
                storage_path, mime_type,
                progress_callback=_quick_progress
            )
            
            extraction_methods = set(p.get("extraction_method", "native") for p in pages_data)
            used_ocr = "ocr" in extraction_methods
            
            # Chunk the document (CPU-only, no DB needed)
            chunks = self._chunk_text_with_pages(pages_data)
            total_chunks = len(chunks)
            
            # ── Phase 3: Embed and store (needs DB, but proceeds quickly) ──
            async with AsyncSessionLocal() as db:
                try:
                    result = await db.execute(
                        select(Document).where(Document.id == document_id)
                    )
                    document = result.scalar_one_or_none()
                    if not document:
                        return
                    
                    # Store extraction metadata
                    document.document_metadata = {
                        **(document.document_metadata or {}),
                        "total_pages": len(pages_data),
                        "document_name": pages_data[0].get("document_name") if pages_data else doc_filename,
                        "extraction_method": "ocr" if used_ocr else "native",
                        "used_ocr": used_ocr,
                    }
                    
                    if total_chunks == 0:
                        document.total_chunks = 0
                        document.total_tokens = 0
                        document.processing_status = "completed"
                        document.processed_at = datetime.now(timezone.utc)
                        await self._update_progress(db, document, "completed", 100, "No content found in document.")
                        return

                    await self._update_progress(db, document, "embedding", 60, 
                        f"Created {total_chunks} chunks. Generating embeddings...")

                    # BATCHED EMBEDDING GENERATION
                    chunk_texts = [chunk_data["text"] for chunk_data in chunks]
                    all_embeddings = [None] * total_chunks
                    
                    for batch_start in range(0, total_chunks, EMBEDDING_BATCH_SIZE):
                        batch_end = min(batch_start + EMBEDDING_BATCH_SIZE, total_chunks)
                        batch_texts = chunk_texts[batch_start:batch_end]
                        
                        batch_embeddings = await self.embedding_service.get_embeddings(batch_texts, is_query=False)
                        
                        for i, emb in enumerate(batch_embeddings):
                            all_embeddings[batch_start + i] = emb
                        
                        embed_progress = 60 + int((batch_end / total_chunks) * 25)
                        await self._update_progress(db, document, "embedding", embed_progress,
                            f"Generating embeddings ({batch_end}/{total_chunks} chunks)...")
                        
                        await asyncio.sleep(0)

                    await self._update_progress(db, document, "storing", 88, 
                        f"Storing {total_chunks} chunks in database...")

                    # BATCHED DB INSERTION
                    total_tokens = 0
                    for batch_start in range(0, total_chunks, DB_INSERT_BATCH_SIZE):
                        batch_end = min(batch_start + DB_INSERT_BATCH_SIZE, total_chunks)
                        chunk_objects = []
                        
                        for idx in range(batch_start, batch_end):
                            chunk_data = chunks[idx]
                            embedding = all_embeddings[idx]
                            
                            chunk_metadata = chunk_data.get("metadata", {})
                            chunk_metadata["source_info"] = {
                                "document_name": chunk_data.get("document_name", doc_filename),
                                "page_number": chunk_data.get("page_number"),
                                "page_range": chunk_data.get("page_range"),
                                "position_in_page": chunk_data.get("position_in_page"),
                                "position_percentage": chunk_data.get("position_percentage"),
                            }
                            
                            chunk = DocumentChunk(
                                document_id=document_id,
                                chunk_index=idx,
                                chunk_text=chunk_data["text"],
                                chunk_embedding=embedding,
                                token_count=chunk_data["token_count"],
                                page_number=chunk_data.get("page_number"),
                                section_heading=chunk_data.get("section_heading"),
                                chunk_metadata=chunk_metadata,
                            )
                            chunk_objects.append(chunk)
                            total_tokens += chunk_data["token_count"]
                        
                        db.add_all(chunk_objects)
                        await db.flush()
                        
                        store_progress = 88 + int(((batch_end) / total_chunks) * 10)
                        await self._update_progress(db, document, "storing", store_progress,
                            f"Stored {batch_end}/{total_chunks} chunks...")

                    # Update document
                    document.total_chunks = total_chunks
                    document.total_tokens = total_tokens
                    document.processing_status = "completed"
                    document.processed_at = datetime.now(timezone.utc)
                    
                    ocr_label = " (OCR)" if (document.document_metadata or {}).get("used_ocr") else ""
                    await self._update_progress(db, document, "completed", 100, 
                        f"Processing complete{ocr_label}. {total_chunks} chunks, {total_tokens} tokens.")
                    
                    logger.info(f"✅ Document '{doc_filename}' processed{ocr_label}: {total_chunks} chunks, {total_tokens} tokens")

                except Exception as e:
                    logger.error(f"❌ Document processing failed for {document_id}: {e}")
                    document.processing_status = "failed"
                    document.document_metadata = {
                        **(document.document_metadata or {}),
                        "error": str(e),
                        "processing_step": "failed",
                        "processing_progress": 0,
                    }
                    await db.commit()
        
        except Exception as e:
            logger.error(f"❌ Document extraction failed for {document_id}: {e}")
            try:
                async with AsyncSessionLocal() as db:
                    result = await db.execute(
                        select(Document).where(Document.id == document_id)
                    )
                    document = result.scalar_one_or_none()
                    if document:
                        document.processing_status = "failed"
                        document.document_metadata = {
                            **(document.document_metadata or {}),
                            "error": str(e),
                            "processing_step": "failed",
                            "processing_progress": 0,
                        }
                        await db.commit()
            except Exception:
                pass

    async def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text from document based on file type (legacy method for compatibility)."""
        pages = await self._extract_text_with_pages(file_path, mime_type)
        return "\n\n".join([p["text"] for p in pages])
    
    async def _extract_text_with_pages(self, file_path: str, mime_type: str, progress_callback=None) -> List[dict]:
        """
        Extract text from document with page-level metadata.
        
        For PDFs, first attempts native text extraction. If the document appears
        to be scanned (sparse text), falls back to OCR with retry logic.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            progress_callback: Optional async callback(step, progress, detail) for progress updates
            
        Returns a list of dicts with structure:
        {
            "page_number": int,
            "text": str,
            "blocks": [{"text": str, "position": {"top": float, "left": float, "bottom": float, "right": float}}]
            "extraction_method": "native" | "ocr"
        }
        """
        import fitz  # PyMuPDF
        from pathlib import Path

        document_name = Path(file_path).stem

        if mime_type == "application/pdf":
            # Step 1: Try native text extraction (run in thread — fitz is CPU-bound)
            pages_data = await asyncio.to_thread(self._extract_pdf_native, file_path, document_name)
            
            # Step 2: Check if the document is scanned/image-heavy
            if settings.OCR_ENABLED and self._is_scanned_pdf(pages_data):
                if progress_callback:
                    await progress_callback(
                        "ocr_detecting", 8,
                        f"Detected scanned/image-heavy PDF ({len(pages_data)} pages). Starting OCR..."
                    )
                
                # Step 3: Fall back to OCR with retry
                ocr_pages = await self._extract_pdf_ocr_with_retry(
                    file_path, document_name, len(pages_data), progress_callback
                )
                
                if ocr_pages:
                    return ocr_pages
                else:
                    # OCR failed completely, return whatever native extraction got
                    if progress_callback:
                        await progress_callback(
                            "ocr_fallback", 14,
                            "OCR failed. Using available native text extraction."
                        )
                    return pages_data
            
            return pages_data
        
        elif mime_type == "text/plain":
            async with aiofiles.open(file_path, "r") as f:
                content = await f.read()
            return [{
                "page_number": 1,
                "text": content,
                "page_height": None,
                "blocks": [{"text": content, "position": {"top": 0, "left": 0, "bottom": 100, "right": 100}}],
                "document_name": document_name,
                "extraction_method": "native",
            }]
        
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            lines_per_page = 40
            pages_data = []
            current_page = 1
            current_blocks = []
            current_line_count = 0
            
            for para in paragraphs:
                para_lines = para.count('\n') + 1
                if current_line_count + para_lines > lines_per_page and current_blocks:
                    pages_data.append({
                        "page_number": current_page,
                        "text": "\n\n".join([b["text"] for b in current_blocks]),
                        "page_height": None,
                        "blocks": current_blocks,
                        "document_name": document_name,
                        "extraction_method": "native",
                    })
                    current_page += 1
                    current_blocks = []
                    current_line_count = 0
                
                current_blocks.append({
                    "text": para,
                    "position": {"top": current_line_count * 2.5, "left": 0, "bottom": (current_line_count + para_lines) * 2.5, "right": 100}
                })
                current_line_count += para_lines
            
            if current_blocks:
                pages_data.append({
                    "page_number": current_page,
                    "text": "\n\n".join([b["text"] for b in current_blocks]),
                    "page_height": None,
                    "blocks": current_blocks,
                    "document_name": document_name,
                    "extraction_method": "native",
                })
            
            return pages_data if pages_data else [{"page_number": 1, "text": "", "page_height": None, "blocks": [], "document_name": document_name, "extraction_method": "native"}]
        
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    def _extract_pdf_native(self, file_path: str, document_name: str) -> List[dict]:
        """Extract text from PDF using PyMuPDF native text extraction."""
        import fitz
        
        pages_data = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                text_blocks = []
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        block_text = ""
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                block_text += span.get("text", "")
                            block_text += "\n"
                        
                        if block_text.strip():
                            text_blocks.append({
                                "text": block_text.strip(),
                                "position": {
                                    "top": block.get("bbox", [0, 0, 0, 0])[1],
                                    "left": block.get("bbox", [0, 0, 0, 0])[0],
                                    "bottom": block.get("bbox", [0, 0, 0, 0])[3],
                                    "right": block.get("bbox", [0, 0, 0, 0])[2],
                                }
                            })
                
                page_height = page.rect.height
                full_text = page.get_text().replace('\x00', '')  # strip null bytes (PG rejects them)
                
                pages_data.append({
                    "page_number": page_num,
                    "text": full_text,
                    "page_height": page_height,
                    "blocks": text_blocks,
                    "document_name": document_name,
                    "extraction_method": "native",
                })
        
        return pages_data

    def _is_scanned_pdf(self, pages_data: List[dict]) -> bool:
        """
        Detect whether a PDF is scanned/image-heavy by analyzing text density.
        
        Uses two heuristics:
        1. Average characters per page is below threshold
        2. Ratio of pages with meaningful text is below threshold
        """
        if not pages_data:
            return False
        
        min_chars = settings.OCR_MIN_TEXT_PER_PAGE
        sparse_threshold = settings.OCR_SPARSE_THRESHOLD
        
        pages_with_text = 0
        total_chars = 0
        
        for page in pages_data:
            text = page.get("text", "").strip()
            char_count = len(text)
            total_chars += char_count
            if char_count >= min_chars:
                pages_with_text += 1
        
        total_pages = len(pages_data)
        text_ratio = pages_with_text / total_pages if total_pages > 0 else 0
        avg_chars_per_page = total_chars / total_pages if total_pages > 0 else 0
        
        # Consider scanned if most pages lack text
        # A normal text PDF page has 1000-3000+ characters
        is_scanned = text_ratio < sparse_threshold or avg_chars_per_page < min_chars
        
        return is_scanned

    async def _extract_pdf_ocr_with_retry(
        self,
        file_path: str,
        document_name: str,
        total_pages: int,
        progress_callback=None,
    ) -> Optional[List[dict]]:
        """
        Extract text from PDF pages using OCR (Tesseract) with retry logic.
        
        Renders each PDF page as an image and runs Tesseract OCR.
        Retries failed pages up to OCR_MAX_RETRIES times.
        
        Returns list of page dicts or None if OCR is unavailable.
        """
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            import pytesseract
            from PIL import Image
            import fitz
        except ImportError as e:
            logger.warning(f"OCR dependencies not available: {e}")
            return None
        
        # Verify tesseract is installed
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            logger.warning("Tesseract binary not found. OCR disabled.")
            return None
        
        max_retries = settings.OCR_MAX_RETRIES
        dpi = settings.OCR_DPI
        language = settings.OCR_LANGUAGE
        
        pages_data = []
        failed_pages = []
        
        if progress_callback:
            await progress_callback(
                "ocr", 10,
                f"Running OCR on {total_pages} pages (this may take a while)..."
            )
        
        try:
            with fitz.open(file_path) as doc:
                for page_idx, page in enumerate(doc):
                    page_num = page_idx + 1
                    ocr_text = None
                    
                    for attempt in range(max_retries + 1):
                        try:
                            # Run CPU-bound rendering + OCR in a thread
                            # so the asyncio event loop stays responsive
                            def _ocr_page(pg=page):
                                zoom = dpi / 72
                                mat = fitz.Matrix(zoom, zoom)
                                pix = pg.get_pixmap(matrix=mat)
                                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                text = pytesseract.image_to_string(
                                    img, lang=language, config='--psm 6'
                                )
                                pix = None
                                img = None
                                return text

                            ocr_text = await asyncio.to_thread(_ocr_page)
                            
                            break  # Success, no need to retry
                            
                        except Exception as e:
                            if attempt < max_retries:
                                logger.warning(
                                    f"OCR attempt {attempt + 1} failed for page {page_num}: {e}. Retrying..."
                                )
                                await asyncio.sleep(0.5 * (attempt + 1))  # Backoff
                            else:
                                logger.error(
                                    f"OCR failed for page {page_num} after {max_retries + 1} attempts: {e}"
                                )
                                failed_pages.append(page_num)
                                ocr_text = ""
                    
                    text = (ocr_text or "").strip()
                    
                    # Build block data from OCR text (paragraph-level)
                    text_blocks = []
                    if text:
                        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                        block_height = page.rect.height / max(len(paragraphs), 1)
                        for i, para in enumerate(paragraphs):
                            text_blocks.append({
                                "text": para,
                                "position": {
                                    "top": i * block_height,
                                    "left": 0,
                                    "bottom": (i + 1) * block_height,
                                    "right": page.rect.width,
                                }
                            })
                    
                    pages_data.append({
                        "page_number": page_num,
                        "text": text,
                        "page_height": page.rect.height,
                        "blocks": text_blocks,
                        "document_name": document_name,
                        "extraction_method": "ocr",
                    })
                    
                    # Update progress periodically
                    if progress_callback and (page_num % 5 == 0 or page_num == total_pages):
                        ocr_progress = 10 + int((page_num / total_pages) * 45)  # 10-55% range
                        await progress_callback(
                            "ocr", min(ocr_progress, 55),
                            f"OCR processing page {page_num}/{total_pages}..."
                        )
            
            if failed_pages:
                logger.warning(f"OCR completed with {len(failed_pages)} failed pages: {failed_pages}")
            
            return pages_data
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return None

    def _chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ) -> List[dict]:
        """
        Chunk text into smaller pieces for embedding using semantic-aware splitting.
        Uses RecursiveCharacterTextSplitter for better context preservation.
        """
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        import re
        
        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP

        # Use RecursiveCharacterTextSplitter for semantic-aware chunking
        # It tries to split on these separators in order, preferring larger semantic units
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n\n",      # Major section breaks
                "\n\n",        # Paragraph breaks
                "\n",          # Line breaks
                ". ",          # Sentence endings
                "? ",          # Question endings
                "! ",          # Exclamation endings
                "; ",          # Semicolon breaks
                ", ",          # Comma breaks
                " ",           # Word breaks
                ""             # Character-level fallback
            ],
            length_function=len,
            is_separator_regex=False,
        )
        
        # Split the text
        raw_chunks = splitter.split_text(text)
        
        # Build chunk metadata
        chunks = []
        for idx, chunk_text in enumerate(raw_chunks):
            chunk_text = chunk_text.strip()
            if not chunk_text:
                continue
            
            # Extract metadata from chunk content
            metadata = self._extract_chunk_metadata(chunk_text)
            
            chunks.append({
                "text": chunk_text,
                "token_count": len(chunk_text.split()),
                "metadata": metadata,
            })

        return chunks

    def _chunk_text_with_pages(
        self,
        pages_data: List[dict],
        chunk_size: int = None,
        chunk_overlap: int = None,
    ) -> List[dict]:
        """
        Chunk text with page-level tracking for source attribution.
        
        Uses an efficient range-based approach instead of per-character dicts
        to keep memory usage low for large documents.
        
        Each chunk contains:
        - text: The chunk content
        - token_count: Number of tokens
        - page_number: Primary page the chunk is from
        - page_range: [start_page, end_page] if chunk spans pages
        - position_in_page: "top", "middle", or "bottom" based on vertical position
        - section_heading: Detected section heading if any
        - document_name: Name of the source document
        - metadata: Additional metadata (has_code, has_math, etc.)
        """
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        import bisect
        
        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n\n", "\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""],
            length_function=len,
            is_separator_regex=False,
        )
        
        chunks = []
        
        # Build a lightweight range-based mapping: (start_char, end_char, page_num, doc_name, page_text_length)
        # This uses O(num_pages) memory instead of O(num_characters)
        full_text = ""
        page_ranges = []  # List of (start_char_idx, page_num, document_name, page_text_length)
        page_starts = []  # Just the start indices for binary search
        
        for page in pages_data:
            page_num = page["page_number"]
            # Strip null bytes — PostgreSQL rejects \x00 in VARCHAR/TEXT columns
            page_text = page["text"].replace('\x00', '')
            document_name = page.get("document_name", "Unknown")
            
            start_idx = len(full_text)
            full_text += page_text + "\n\n"
            
            page_ranges.append((start_idx, page_num, document_name, len(page_text)))
            page_starts.append(start_idx)
        
        def _get_page_info_at(char_idx: int):
            """Get page info for a character index using binary search. O(log n)."""
            idx = bisect.bisect_right(page_starts, char_idx) - 1
            if idx < 0:
                idx = 0
            start, page_num, doc_name, page_len = page_ranges[idx]
            char_in_page = char_idx - start
            return page_num, doc_name, char_in_page, page_len
        
        # Split the combined text
        raw_chunks = splitter.split_text(full_text)
        
        # Map each chunk back to its source pages using range lookups (not per-char)
        current_pos = 0
        for chunk_text in raw_chunks:
            chunk_text_stripped = chunk_text.strip()
            if not chunk_text_stripped:
                continue
            
            # Find where this chunk starts in the full text
            chunk_start = full_text.find(chunk_text, current_pos)
            if chunk_start == -1:
                chunk_start = current_pos
            chunk_end = chunk_start + len(chunk_text)
            current_pos = chunk_start + 1
            
            # Determine page range using start/end + midpoint sampling
            start_page, start_doc, start_char_in_page, start_page_len = _get_page_info_at(chunk_start)
            end_page, end_doc, end_char_in_page, end_page_len = _get_page_info_at(min(chunk_end - 1, len(full_text) - 1))
            mid_page, _, mid_char_in_page, mid_page_len = _get_page_info_at((chunk_start + chunk_end) // 2)
            
            primary_page = start_page
            page_range = [start_page, end_page]
            document_name = start_doc
            
            # Estimate position in page from start position
            if start_page_len > 0:
                avg_position = (start_char_in_page / start_page_len) * 100
            else:
                avg_position = 50.0
            
            if avg_position < 33:
                position_label = "top"
            elif avg_position < 66:
                position_label = "middle"
            else:
                position_label = "bottom"
            
            # Extract section heading if present
            section_heading = self._extract_section_heading(chunk_text_stripped)
            
            # Extract metadata
            metadata = self._extract_chunk_metadata(chunk_text_stripped)
            metadata["source_info"] = {
                "page_number": primary_page,
                "page_range": page_range,
                "position_in_page": position_label,
                "position_percentage": round(avg_position, 1),
                "document_name": document_name,
            }
            
            chunks.append({
                "text": chunk_text_stripped,
                "token_count": len(chunk_text_stripped.split()),
                "page_number": primary_page,
                "page_range": page_range,
                "position_in_page": position_label,
                "position_percentage": round(avg_position, 1),
                "section_heading": section_heading,
                "document_name": document_name,
                "metadata": metadata,
            })
        
        return chunks
    
    def _extract_section_heading(self, text: str) -> Optional[str]:
        """Extract section heading from chunk if present at the beginning."""
        import re
        
        lines = text.strip().split('\n')
        if not lines:
            return None
        
        first_line = lines[0].strip()
        
        # Common heading patterns
        heading_patterns = [
            r'^(?:Chapter|Section|Part|Unit)\s+\d+[.:]?\s*(.+)$',
            r'^(\d+\.)+\s*(.+)$',  # Numbered headings like "1.2.3 Title"
            r'^([A-Z][A-Za-z\s]+):$',  # "Title:" format
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS headings
        ]
        
        for pattern in heading_patterns:
            match = re.match(pattern, first_line, re.IGNORECASE)
            if match:
                # Return the heading (last group usually contains the title)
                return match.group(match.lastindex) if match.lastindex else first_line
        
        # If first line is short and looks like a title (< 80 chars, no sentence-ending punctuation)
        if len(first_line) < 80 and not re.search(r'[.!?]$', first_line):
            return first_line
        
        return None

    def _extract_chunk_metadata(self, text: str) -> dict:
        """Extract metadata from chunk content for filtering."""
        import re
        
        metadata = {}
        
        # Check for code blocks
        metadata["has_code"] = bool(re.search(
            r'```|def\s+\w+|class\s+\w+|function\s+\w+|import\s+\w+|from\s+\w+|const\s+\w+|let\s+\w+|var\s+\w+',
            text
        ))
        
        # Check for mathematical content
        metadata["has_math"] = bool(re.search(
            r'\$.*\$|\d+\s*[+\-*/=<>≤≥]\s*\d+|√|∑|∏|∫|∂|∞|π|θ|α|β|γ|equation|formula|theorem',
            text,
            re.IGNORECASE
        ))
        
        # Check for lists
        metadata["has_list"] = bool(re.search(
            r'^\s*[-•*]\s+|^\s*\d+\.\s+|^\s*[a-z]\)\s+|^\s*[ivxlcdm]+\.\s+',
            text,
            re.MULTILINE | re.IGNORECASE
        ))
        
        # Check for tables
        metadata["has_table"] = bool(re.search(
            r'\|.*\|.*\||^\s*\+[-+]+\+\s*$',
            text,
            re.MULTILINE
        ))
        
        # Check for definitions/key terms
        metadata["has_definition"] = bool(re.search(
            r':\s*definition|defined\s+as|refers\s+to|means\s+that|is\s+known\s+as',
            text,
            re.IGNORECASE
        ))
        
        # Estimate complexity/difficulty based on various heuristics
        metadata["estimated_complexity"] = self._estimate_complexity(text)
        
        return metadata

    def _estimate_complexity(self, text: str) -> str:
        """Estimate the complexity of text content."""
        import re
        
        # Count complexity indicators
        complexity_score = 0
        
        # Long words (likely technical terms)
        long_words = len([w for w in text.split() if len(w) > 10])
        complexity_score += min(long_words / 5, 2)
        
        # Sentences per 100 words (lower = more complex)
        words = len(text.split())
        sentences = len(re.findall(r'[.!?]+', text))
        if words > 0 and sentences > 0:
            avg_sentence_length = words / sentences
            if avg_sentence_length > 25:
                complexity_score += 1
            if avg_sentence_length > 35:
                complexity_score += 1
        
        # Technical patterns
        if re.search(r'theorem|proof|lemma|corollary|hypothesis', text, re.IGNORECASE):
            complexity_score += 2
        if re.search(r'algorithm|complexity|optimization', text, re.IGNORECASE):
            complexity_score += 1
        
        # Classify
        if complexity_score >= 4:
            return "high"
        elif complexity_score >= 2:
            return "medium"
        else:
            return "low"

    async def get_document(self, document_id: uuid.UUID, user_id: uuid.UUID) -> Optional[Document]:
        """Get document by ID (user-scoped)."""
        result = await self.db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )
        return result.scalar_one_or_none()

    async def get_documents(
        self,
        user_id: uuid.UUID,
        page: int = 1,
        limit: int = 20,
        status: Optional[str] = None,
    ) -> tuple[List[Document], PaginationInfo]:
        """Get paginated documents for a user."""
        query = select(Document).where(Document.user_id == user_id)
        
        if status:
            query = query.where(Document.processing_status == status)
        
        # Count total
        count_query = select(func.count()).select_from(query.subquery())
        total_result = await self.db.execute(count_query)
        total = total_result.scalar_one()

        # Get page
        offset = (page - 1) * limit
        query = query.offset(offset).limit(limit).order_by(Document.upload_timestamp.desc())
        result = await self.db.execute(query)
        documents = result.scalars().all()

        # Get question counts
        for doc in documents:
            q_result = await self.db.execute(
                select(func.count()).where(Question.document_id == doc.id)
            )
            doc.questions_generated = q_result.scalar_one()

        pagination = PaginationInfo(
            page=page,
            limit=limit,
            total=total,
            total_pages=(total + limit - 1) // limit,
        )

        return documents, pagination

    async def delete_document(self, document_id: uuid.UUID, user_id: uuid.UUID) -> bool:
        """Delete a document and all related data."""
        document = await self.get_document(document_id, user_id)
        if not document:
            return False

        # Delete file only if no other document rows share the same storage path.
        other_ref_count_result = await self.db.execute(
            select(func.count()).where(
                Document.storage_path == document.storage_path,
                Document.id != document.id,
            )
        )
        other_ref_count = int(other_ref_count_result.scalar_one() or 0)
        if other_ref_count == 0 and os.path.exists(document.storage_path):
            os.remove(document.storage_path)

        # Delete from database (cascades to chunks and questions)
        await self.db.delete(document)
        await self.db.commit()
        return True

    async def get_document_chunks(
        self,
        document_id: uuid.UUID,
        user_id: uuid.UUID,
    ) -> List[DocumentChunk]:
        """Get all chunks for a document."""
        # Verify ownership
        document = await self.get_document(document_id, user_id)
        if not document:
            raise ValueError("Document not found")

        result = await self.db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index)
        )
        return result.scalars().all()

    async def search_chunks(
        self,
        document_id: uuid.UUID,
        query_embedding: List[float],
        top_k: int = 5,
    ) -> List[DocumentChunk]:
        """Search chunks by semantic similarity."""
        from pgvector.sqlalchemy import Vector
        
        result = await self.db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_embedding.cosine_distance(query_embedding))
            .limit(top_k)
        )
        return result.scalars().all()

    @staticmethod
    def _tokenize_for_search(text: str) -> List[str]:
        """Tokenize text for robust keyword/BM25 scoring."""
        import re

        stopwords = {
            "the", "a", "an", "and", "or", "but", "if", "then", "else", "when", "while",
            "for", "to", "of", "in", "on", "at", "by", "from", "with", "as", "is", "are",
            "was", "were", "be", "been", "being", "this", "that", "these", "those", "it", "its",
            "which", "what", "who", "whom", "whose", "why", "how", "where", "do", "does", "did",
            "can", "could", "would", "should", "will", "shall", "about", "into", "through",
        }

        tokens = re.findall(r"[a-zA-Z0-9_]+", (text or "").lower())
        filtered = [t for t in tokens if len(t) > 1 and t not in stopwords]
        return filtered

    @staticmethod
    def _keyword_overlap_ratio(query_terms: List[str], text_terms: List[str]) -> float:
        """Compute keyword overlap ratio between query terms and chunk terms."""
        if not query_terms:
            return 0.0
        query_set = set(query_terms)
        if not query_set:
            return 0.0
        text_set = set(text_terms)
        return len(query_set.intersection(text_set)) / len(query_set)

    async def hybrid_search(
        self,
        document_id: uuid.UUID,
        query: str,
        query_embedding: List[float],
        top_k: int = 5,
        alpha: float = 0.5,
    ) -> List[DocumentChunk]:
        """
        Hybrid search combining vector similarity and BM25 keyword matching.
        
        Args:
            document_id: Document to search within
            query: The text query for BM25 scoring
            query_embedding: Pre-computed embedding for semantic search
            top_k: Number of results to return
            alpha: Balance between semantic (1.0) and keyword (0.0) search
        
        Returns:
            Top-k chunks ranked by combined score
        """
        from rank_bm25 import BM25Okapi
        
        # Get all chunks for this document
        result = await self.db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = list(result.scalars().all())
        
        if not chunks:
            return []
        
        query_terms = self._tokenize_for_search(query)
        if not query_terms:
            query_terms = query.lower().split()

        # BM25 scoring with cleaner tokenization
        tokenized_chunks = [self._tokenize_for_search(c.chunk_text) for c in chunks]
        tokenized_chunks = [toks if toks else c.chunk_text.lower().split() for c, toks in zip(chunks, tokenized_chunks)]
        bm25 = BM25Okapi(tokenized_chunks)
        bm25_scores = bm25.get_scores(query_terms)
        
        # Normalize BM25 scores to [0, 1]
        max_bm25 = max(bm25_scores) if max(bm25_scores) > 0 else 1
        bm25_scores_normalized = [s / max_bm25 for s in bm25_scores]
        
        # Vector similarity scores
        vector_scores = []
        for chunk in chunks:
            # Check if embedding exists (handle numpy array truth value issue)
            if chunk.chunk_embedding is not None and len(chunk.chunk_embedding) > 0:
                sim = self.embedding_service.compute_similarity(
                    query_embedding, chunk.chunk_embedding
                )
                # Convert to [0, 1] range (cosine similarity is already [-1, 1])
                vector_scores.append((sim + 1) / 2)
            else:
                vector_scores.append(0)
        
        # Keyword overlap score (guards against semantically-close but off-topic chunks)
        overlap_scores = [
            self._keyword_overlap_ratio(query_terms, tokenized_chunks[i])
            for i in range(len(chunks))
        ]

        # Combine scores using weighted sum
        lexical_weight = (1 - alpha)
        bm25_weight = lexical_weight * 0.7
        overlap_weight = lexical_weight * 0.3
        combined_scores = [
            (alpha * vector_scores[i])
            + (bm25_weight * bm25_scores_normalized[i])
            + (overlap_weight * overlap_scores[i])
            for i in range(len(chunks))
        ]

        # Penalize low-signal matches to reduce irrelevant retrieval.
        for i in range(len(chunks)):
            if vector_scores[i] < 0.35 and overlap_scores[i] < 0.08:
                combined_scores[i] *= 0.65
        
        # Sort by combined score and return top-k
        ranked = sorted(
            zip(chunks, combined_scores),
            key=lambda x: x[1],
            reverse=True
        )
        
        return [chunk for chunk, _ in ranked[:top_k]]

    async def hybrid_search_multi_document(
        self,
        document_ids: List[uuid.UUID],
        query: str,
        query_embedding: List[float],
        top_k: int = 5,
        alpha: float = 0.5,
    ) -> List[DocumentChunk]:
        """
        Hybrid search combining vector similarity and BM25 keyword matching
        across multiple documents (e.g. primary + all reference books).
        
        Args:
            document_ids: List of document IDs to search across
            query: The text query for BM25 scoring
            query_embedding: Pre-computed embedding for semantic search
            top_k: Number of results to return
            alpha: Balance between semantic (1.0) and keyword (0.0) search
        
        Returns:
            Top-k chunks ranked by combined score across all documents
        """
        from rank_bm25 import BM25Okapi
        
        if not document_ids:
            return []
        
        # Get all chunks across all documents
        result = await self.db.execute(
            select(DocumentChunk)
            .options(selectinload(DocumentChunk.document))
            .where(DocumentChunk.document_id.in_(document_ids))
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = list(result.scalars().all())
        
        if not chunks:
            return []
        
        query_terms = self._tokenize_for_search(query)
        if not query_terms:
            query_terms = query.lower().split()

        # BM25 scoring with cleaner tokenization
        tokenized_chunks = [self._tokenize_for_search(c.chunk_text) for c in chunks]
        tokenized_chunks = [toks if toks else c.chunk_text.lower().split() for c, toks in zip(chunks, tokenized_chunks)]
        bm25 = BM25Okapi(tokenized_chunks)
        bm25_scores = bm25.get_scores(query_terms)
        
        # Normalize BM25 scores to [0, 1]
        max_bm25 = max(bm25_scores) if max(bm25_scores) > 0 else 1
        bm25_scores_normalized = [s / max_bm25 for s in bm25_scores]
        
        # Vector similarity scores
        vector_scores = []
        for chunk in chunks:
            if chunk.chunk_embedding is not None and len(chunk.chunk_embedding) > 0:
                sim = self.embedding_service.compute_similarity(
                    query_embedding, chunk.chunk_embedding
                )
                vector_scores.append((sim + 1) / 2)
            else:
                vector_scores.append(0)
        
        # Keyword overlap score
        overlap_scores = [
            self._keyword_overlap_ratio(query_terms, tokenized_chunks[i])
            for i in range(len(chunks))
        ]

        # Combine scores using weighted sum
        lexical_weight = (1 - alpha)
        bm25_weight = lexical_weight * 0.7
        overlap_weight = lexical_weight * 0.3
        combined_scores = [
            (alpha * vector_scores[i])
            + (bm25_weight * bm25_scores_normalized[i])
            + (overlap_weight * overlap_scores[i])
            for i in range(len(chunks))
        ]

        # Penalize low-signal matches and softly prefer primary chunks when present.
        for i, chunk in enumerate(chunks):
            if vector_scores[i] < 0.35 and overlap_scores[i] < 0.08:
                combined_scores[i] *= 0.6

            doc = getattr(chunk, "document", None)
            if doc and getattr(doc, "index_type", None) == "primary":
                combined_scores[i] *= 1.06
        
        # Sort by combined score
        ranked = sorted(
            zip(chunks, combined_scores),
            key=lambda x: x[1],
            reverse=True
        )

        # Balance across documents to avoid one noisy reference doc dominating results.
        max_per_document = max(2, top_k // 2)
        selected: List[DocumentChunk] = []
        doc_counts: Dict[uuid.UUID, int] = {}

        for chunk, _ in ranked:
            doc_count = doc_counts.get(chunk.document_id, 0)
            if doc_count >= max_per_document:
                continue
            selected.append(chunk)
            doc_counts[chunk.document_id] = doc_count + 1
            if len(selected) >= top_k:
                break

        if len(selected) < top_k:
            selected_ids = {c.id for c in selected}
            for chunk, _ in ranked:
                if chunk.id in selected_ids:
                    continue
                selected.append(chunk)
                if len(selected) >= top_k:
                    break

        primary_candidates = [c for c, _ in ranked if getattr(getattr(c, "document", None), "index_type", None) == "primary"]
        if primary_candidates and selected and all(
            getattr(getattr(c, "document", None), "index_type", None) != "primary" for c in selected
        ):
            selected[-1] = primary_candidates[0]

        return selected[:top_k]

    async def get_reference_document_ids(
        self,
        user_id: uuid.UUID,
        subject_id: uuid.UUID,
        index_type: str = "reference_book",
    ) -> List[uuid.UUID]:
        """
        Get IDs of all completed reference documents for a subject.
        Used to include reference book chunks in question generation searches.
        """
        result = await self.db.execute(
            select(Document.id).where(
                Document.user_id == user_id,
                Document.subject_id == subject_id,
                Document.index_type == index_type,
                Document.processing_status == "completed",
            )
        )
        return list(result.scalars().all())

    async def get_reference_documents(
        self,
        user_id: uuid.UUID,
        subject_id: Optional[uuid.UUID] = None,
        index_type: Optional[str] = None,
    ) -> List[Document]:
        """
        Get reference documents (books, template papers, and reference questions) for a user.
        Includes all processing statuses so users can see pending/processing uploads.
        """
        query = select(Document).where(
            Document.user_id == user_id,
            Document.index_type.in_(["reference_book", "template_paper", "reference_questions"]),
        )
        
        if subject_id:
            query = query.where(Document.subject_id == subject_id)
        
        if index_type:
            query = query.where(Document.index_type == index_type)
        
        query = query.order_by(Document.upload_timestamp.desc())
        
        result = await self.db.execute(query)
        return result.scalars().all()

    async def get_reference_chunks(
        self,
        user_id: uuid.UUID,
        subject_id: Optional[uuid.UUID] = None,
        index_type: Optional[str] = None,
        query_embedding: Optional[List[float]] = None,
        top_k: int = 10,
    ) -> List[DocumentChunk]:
        """
        Get chunks from reference documents.
        If query_embedding is provided, returns the most relevant chunks.
        Otherwise returns all chunks.
        """
        # Get reference documents
        docs = await self.get_reference_documents(user_id, subject_id, index_type)
        
        if not docs:
            return []
        
        doc_ids = [doc.id for doc in docs]
        
        # Get all chunks with document relationship for filename access
        result = await self.db.execute(
            select(DocumentChunk)
            .options(selectinload(DocumentChunk.document))
            .where(DocumentChunk.document_id.in_(doc_ids))
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = result.scalars().all()
        
        if not chunks:
            return []
        
        # If no query embedding, return all chunks
        if query_embedding is None:
            return chunks
        
        # Score and rank by similarity
        scored_chunks = []
        for chunk in chunks:
            if chunk.chunk_embedding is not None and len(chunk.chunk_embedding) > 0:
                similarity = self.embedding_service.compute_similarity(
                    query_embedding, chunk.chunk_embedding
                )
                scored_chunks.append((chunk, similarity))
        
        # Sort by similarity
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        
        return [c for c, _ in scored_chunks[:top_k]]

    async def get_primary_chunks_excluding_used(
        self,
        document_id: uuid.UUID,
        used_chunk_ids: List[uuid.UUID],
        query_embedding: Optional[List[float]] = None,
        top_k: int = 5,
    ) -> List[DocumentChunk]:
        """
        Get primary document chunks excluding already-used ones.
        Useful for intelligent regeneration with alternative chunks.
        """
        # Get all chunks for the document
        result = await self.db.execute(
            select(DocumentChunk)
            .where(
                DocumentChunk.document_id == document_id,
                ~DocumentChunk.id.in_(used_chunk_ids) if used_chunk_ids else True,
            )
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = result.scalars().all()
        
        if not chunks:
            return []
        
        # If no query embedding, return random selection
        if query_embedding is None:
            import random
            return random.sample(chunks, min(top_k, len(chunks)))
        
        # Score and rank by similarity
        scored_chunks = []
        for chunk in chunks:
            if chunk.chunk_embedding is not None and len(chunk.chunk_embedding) > 0:
                similarity = self.embedding_service.compute_similarity(
                    query_embedding, chunk.chunk_embedding
                )
                scored_chunks.append((chunk, similarity))
        
        # Sort by similarity
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        
        return [c for c, _ in scored_chunks[:top_k]]

    async def get_reference_questions(
        self,
        user_id: uuid.UUID,
        subject_id: uuid.UUID,
        question_type: Optional[str] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Get parsed reference questions for a subject.
        These are used as style/standard templates for question generation.
        
        Reference questions are stored as chunks in documents with index_type='reference_questions'.
        Each chunk contains a JSON-encoded question with fields like:
        - question_text, options, correct_answer, question_type, difficulty, marks, co, lo
        """
        import json
        
        # Get reference question document IDs
        doc_ids = await self.get_reference_document_ids(
            user_id=user_id,
            subject_id=subject_id,
            index_type="reference_questions",
        )
        
        if not doc_ids:
            return []
        
        # Get chunks from these documents
        result = await self.db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id.in_(doc_ids))
            .order_by(DocumentChunk.chunk_index)
            .limit(limit * 3)  # Get more to filter by type
        )
        chunks = result.scalars().all()
        
        parsed_questions = []
        for chunk in chunks:
            try:
                # Reference question chunks store JSON in chunk_text
                q_data = json.loads(chunk.chunk_text)
                if isinstance(q_data, dict) and q_data.get("question_text"):
                    # Filter by type if specified
                    if question_type and q_data.get("question_type") != question_type:
                        continue
                    parsed_questions.append(q_data)
                    if len(parsed_questions) >= limit:
                        break
            except (json.JSONDecodeError, TypeError):
                # Chunk might be plain text, skip it
                continue
        
        return parsed_questions
